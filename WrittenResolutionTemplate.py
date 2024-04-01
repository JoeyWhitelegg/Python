# Resolution in lieu of a meeting template

# This program will write up a template for a typical resolution written by a company
# Will prompt user to fill in the required info for each section in the resolution

#! python3

import docx
import datetime as dt
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx import Document


def creation(res_doc):
# SET TO LEGAL SPECS
    # set font style
    font_styles = res_doc.styles
    font_charstyle = font_styles.add_style('JoeyStyle', WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.name = 'Times New Roman'
    font_object.size = Pt(12)

    # margins
    sections = res_doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.right_margin = Inches(1)
        section.left_margin = Inches(1)

    # set up paragraph spacing and line spacing
    # line spacing = single
    # paragraph spacing = 12 before

    # formatting, done individual, want a way to do a uniform font and size for entire doc
    print('Please enter the business name.')
    input_name = input()
    bus_name = input_name.upper()

    # resolutions from business at top justification = center
    res_title = res_doc.add_paragraph()
    res_title.add_run('RESOLUTIONS OF ', style='JoeyStyle')
    bus_namerun = res_title.add_run(bus_name, style='JoeyStyle')
    bus_namerun.bold = True
    res_title.add_run('\nPREPARED IN WRITING ON', style='JoeyStyle')
    restit_format = res_title.paragraph_format
    restit_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # date justified center
    today = dt.date.today()
    date = (f'{today:%B %d, %Y}')
    res_title_date = res_doc.add_paragraph()
    res_title_date.add_run(date + '.', style='JoeyStyle')
    date_format = res_title_date.paragraph_format
    date_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # the why paragraphs justified to the left with font and point size set
    print('Type in the resolution.')
    whereas_one = input()
    font_w = ('WHEREAS ')
    wpara_one = res_doc.add_paragraph()
    font_wrun = wpara_one.add_run(font_w, style='JoeyStyle')
    font_wrun.bold = True
    wpara_one.add_run((str(whereas_one)), style='JoeyStyle')
    wpone_format = wpara_one.paragraph_format
    wpone_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print('Type in any other aspects of resolution.')
    whereas_two = input()
    font_aw = ('AND WHEREAS ')
    wpara_two = res_doc.add_paragraph()
    font_awrun = wpara_two.add_run(font_aw, style='JoeyStyle')
    font_awrun.bold = True
    wpara_two.add_run(str(whereas_two), style='JoeyStyle')
    wptwo_format = wpara_two.paragraph_format
    wptwo_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # add any extra paragraphs needed to talk about the resolution
    while True:
        print('Type in any other aspects for resolution or type nothing to continue.')
        whereas_three = input()
        if whereas_three == '':
            break
        font_aww = ('AND WHEREAS ')
        wpara_three = res_doc.add_paragraph()
        font_awwrun = wpara_three.add_run(font_aww, style='JoeyStyle')
        font_awwrun.bold = True
        wpthree_run = wpara_three.add_run(str(whereas_three), style='JoeyStyle')
        wpthree_format = wpara_three.paragraph_format
        wpthree_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    font_now = ('NOW THEREFORE IT BE RESOLVED THAT:')
    now_para = res_doc.add_paragraph()
    font_nowrun = now_para.add_run(font_now, style='JoeyStyle')
    font_nowrun.bold = True
    nowp_format = now_para.paragraph_format
    nowp_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print('Type in what is the process of the resolution.')
    resolve_one = input()
    rpara_one = res_doc.add_paragraph()
    rpara_one.add_run(('1.\t ' + str(resolve_one)), style='JoeyStyle')

    # any extra processes or not
    extra_proc = []
    while True:
        print('Enter anything extra to add to the process of the resolution or type nothing to continue.')
        reso_two = input()
        if reso_two == '':
            break
        extra_proc = extra_proc + [reso_two]
        number_count = (int(len(extra_proc)) + 1)
        resolve_two = (str(number_count) + ' \t ') + reso_two
        rpara_two = res_doc.add_paragraph()
        rpara_two.add_run(resolve_two, style='JoeyStyle')

    print('Type in when this resolution will go into effect.')
    resolve_inp = input()
    number_count = (int(len(extra_proc)) + 1)
    resolve_fin = str(number_count + 1) + ' \t ' + resolve_inp
    rpara_fin = res_doc.add_paragraph()
    rpara_fin.add_run(resolve_fin, style='JoeyStyle')

    # consented on
    consent_para = res_doc.add_paragraph()
    consent_para.add_run('Consented to on ' + date + '.' + '\n\n', style='JoeyStyle')
    consent_format = consent_para.paragraph_format
    consent_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # add optional amount of signatures
    line = '_' * 50
    def sig_list(signatures):
        ind_name = ''
        for name in range(len(signatures) - 1):
            ind_name += str(signatures[name]) + line + '\n\n\n'
        ind_name += str(signatures[len(signatures) - 1]) + line
        res_doc.add_paragraph(ind_name)

    signatures = []
    while True:
        print('Enter the name of director/shareholder ' + str(len(signatures) + 1) + ' who will be signing\
 this resolution.' + '(Or enter nothing to finish.)')
        sig_name = input()
        if sig_name == '':
            break
        signatures += [sig_name]

    sig_list(signatures)

# create a word file
res_doc = docx.Document()
creation(res_doc)

# save file to location
def sav_doc(res_doc):
    print('Please print initials.')
    initials = input()
    r_init = initials.upper()
    print('Please enter resolution document number.')
    r_number = input()
    init_num = initials + ' - ' + str(r_number)
    res_doc.save('Resolution - ' + init_num + '.docx')

sav_doc(res_doc)
